import io
import os
from dotenv import load_dotenv
from werkzeug.datastructures import FileStorage

# Load environment variables and set up debug functionality
load_dotenv()
DEBUG = os.getenv("DEBUG", "False").lower() in ["true", "1", "t", "yes", "y"]

def debug_log(message):
    if DEBUG:
        print(f"DEBUG [Word Processor]: {message}")

def process_word(file_obj):
    """
    Process a Word document (.docx) and extract its text content.
    
    Args:
        file_obj: File object from request.files
        
    Returns:
        str: Extracted text content in markdown format
    """
    debug_log(f"Processing Word file: {getattr(file_obj, 'filename', 'Unnamed Document')}")
    
    try:
        # Import docx package here to avoid dependency issues if not installed
        try:
            import docx
        except ImportError:
            debug_log("python-docx package not installed. Installing...")
            import subprocess
            subprocess.check_call(["pip", "install", "python-docx"])
            import docx
            debug_log("python-docx package installed successfully")
        
        # Get the file content
        if isinstance(file_obj, FileStorage):
            docx_bytes = file_obj.read()
            file_obj.close()
            docx_stream = io.BytesIO(docx_bytes)
        else:
            # If it's already a bytes object or BytesIO
            docx_stream = file_obj if isinstance(file_obj, io.BytesIO) else io.BytesIO(file_obj)
        
        # Open the Word document
        doc = docx.Document(docx_stream)
        debug_log(f"Word document opened successfully with {len(doc.paragraphs)} paragraphs")
        
        # Initialize document content
        content = f"# Document: {getattr(file_obj, 'filename', 'Unnamed Document')}\n\n"
        
        # Extract text from paragraphs, preserving structure
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                # Check if the paragraph is a heading based on its style
                if para.style.name.startswith('Heading'):
                    heading_level = int(para.style.name.split(' ')[-1]) if para.style.name[-1].isdigit() else 2
                    heading_marks = '#' * min(heading_level, 6)  # Markdown supports heading levels 1-6
                    content += f"{heading_marks} {para.text.strip()}\n\n"
                else:
                    content += f"{para.text.strip()}\n\n"
        
        # Extract tables if any and convert them to markdown tables
        for i, table in enumerate(doc.tables):
            content += f"\n**Table {i+1}**\n\n"
            
            # Create markdown table
            rows = []
            for j, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            
            if rows:
                # Create header row
                header_row = rows[0]
                content += "| " + " | ".join(header_row) + " |\n"
                content += "| " + " | ".join(["---" for _ in header_row]) + " |\n"
                
                # Add data rows
                for row in rows[1:]:
                    # Pad shorter rows to match header length
                    row = row + [""] * (len(header_row) - len(row))
                    content += "| " + " | ".join(row) + " |\n"
                
                content += "\n"
                
        debug_log(f"Word document processing complete: {len(content)} characters extracted")
        return content.strip()
    except Exception as e:
        debug_log(f"ERROR processing Word document: {str(e)}")
        return f"Error processing Word document: {str(e)}"
